#!/usr/bin/env python3
"""
ToMamdas Encoder - Hackathon Challenge Solution
Encodes a tar file into multiple Word documents with 1x1 tables containing base64 chunks.
"""

import os
import sys
import base64
import hashlib
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
CHUNK_SIZE = 30 * 1024 * 1024  # 30MB per chunk (base64 text) - creates ~10 files for intelligraph.tar
OUTPUT_PREFIX = "chunk_"
OUTPUT_EXTENSION = ".docx"

def calculate_sha256(data):
    """Calculate SHA256 hash of data."""
    return hashlib.sha256(data).hexdigest()

def create_word_with_table(chunk_data, chunk_num, total_chunks, chunk_hash):
    """Create a Word document with a 1x1 table containing the chunk data."""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run(f"ToMamdas Chunk {chunk_num:03d} of {total_chunks:03d}")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata paragraph
    metadata = doc.add_paragraph()
    metadata.add_run(f"Chunk: {chunk_num}/{total_chunks}\n")
    metadata.add_run(f"SHA256: {chunk_hash}\n")
    metadata.add_run(f"Size: {len(chunk_data)} bytes")
    metadata_format = metadata.paragraph_format
    metadata_format.space_after = Pt(12)
    
    # Create 1x1 table
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    
    # Set table width to full page width
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(6.5)
    
    # Add chunk data to the single cell
    cell = table.rows[0].cells[0]
    cell.text = chunk_data
    
    # Set font to monospace for better readability
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
    
    return doc

def encode_tar_to_word(tar_path, output_dir="tomamdas_output"):
    """
    Main encoding function: reads tar file, splits into chunks, creates Word documents.
    """
    print(f"ToMamdas Encoder - Starting encoding process...")
    print(f"Input file: {tar_path}")
    
    # Check if input file exists
    if not os.path.exists(tar_path):
        print(f"ERROR: File '{tar_path}' not found!")
        return False
    
    # Get file size
    file_size = os.path.getsize(tar_path)
    print(f"File size: {file_size:,} bytes ({file_size / (1024*1024):.2f} MB)")
    
    # Create output directory
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Created output directory: {output_dir}")
    
    # Read the tar file in binary mode
    print("Reading tar file...")
    with open(tar_path, 'rb') as f:
        binary_data = f.read()
    
    # Calculate original file hash
    original_hash = calculate_sha256(binary_data)
    print(f"Original file SHA256: {original_hash}")
    
    # Convert to base64
    print("Converting to base64...")
    base64_data = base64.b64encode(binary_data).decode('ascii')
    base64_size = len(base64_data)
    print(f"Base64 size: {base64_size:,} bytes ({base64_size / (1024*1024):.2f} MB)")
    
    # Calculate number of chunks
    total_chunks = (base64_size + CHUNK_SIZE - 1) // CHUNK_SIZE
    print(f"Will create {total_chunks} Word document(s)")
    
    # Split into chunks and create Word documents
    print("\nCreating Word documents...")
    for i in range(total_chunks):
        chunk_num = i + 1
        start_idx = i * CHUNK_SIZE
        end_idx = min(start_idx + CHUNK_SIZE, base64_size)
        chunk_data = base64_data[start_idx:end_idx]
        
        # Calculate chunk hash
        chunk_hash = calculate_sha256(chunk_data.encode('ascii'))
        
        # Create Word document
        doc = create_word_with_table(chunk_data, chunk_num, total_chunks, chunk_hash)
        
        # Save document
        output_filename = f"{OUTPUT_PREFIX}{chunk_num:03d}{OUTPUT_EXTENSION}"
        output_path = os.path.join(output_dir, output_filename)
        doc.save(output_path)
        
        print(f"  [{chunk_num}/{total_chunks}] Created {output_filename} ({len(chunk_data):,} bytes)")
    
    # Save metadata file
    metadata_path = os.path.join(output_dir, "metadata.txt")
    with open(metadata_path, 'w') as f:
        f.write(f"ToMamdas Encoding Metadata\n")
        f.write(f"=" * 50 + "\n")
        f.write(f"Original file: {os.path.basename(tar_path)}\n")
        f.write(f"Original size: {file_size} bytes\n")
        f.write(f"Original SHA256: {original_hash}\n")
        f.write(f"Base64 size: {base64_size} bytes\n")
        f.write(f"Total chunks: {total_chunks}\n")
        f.write(f"Chunk size: {CHUNK_SIZE} bytes\n")
    
    print(f"\nMetadata saved to: {metadata_path}")
    print(f"\n[OK] Encoding complete!")
    print(f"[OK] Created {total_chunks} Word document(s) in '{output_dir}' directory")
    print(f"[OK] Original file hash: {original_hash}")
    
    return True

def main():
    """Main entry point."""
    if len(sys.argv) < 2:
        print("Usage: python ToMamdas_encode.py <tar_file> [output_directory]")
        print("\nExample:")
        print("  python ToMamdas_encode.py intelligraph.tar")
        print("  python ToMamdas_encode.py intelligraph.tar my_output_folder")
        sys.exit(1)
    
    tar_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else "tomamdas_output"
    
    try:
        success = encode_tar_to_word(tar_path, output_dir)
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"\nERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()

# Made with Bob
